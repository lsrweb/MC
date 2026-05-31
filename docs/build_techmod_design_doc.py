from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "docs/TechMod_Design.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)
    paragraph.paragraph_format.space_after = Pt(2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return paragraph


def add_heading(doc, text, level):
    paragraph = doc.add_heading(level=level)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.color.rgb = RGBColor(46, 116, 181) if level < 3 else RGBColor(31, 77, 120)
    return paragraph


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Tech Mod 科技模组设计文档")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("面向 Minecraft Forge 1.20.1 的工业科技、能源、材料与自动化路线")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(85, 85, 85)

    add_heading(doc, "1. 设计目标", 1)
    add_paragraph(doc, "Tech Mod 的目标不是一次性堆出大量空注册内容，而是把每个新增内容做成可使用、可扩展、资源完整、行为明确的模块。每个方块和机器都必须拥有独立类、模型、贴图、语言、配方和测试路径。")
    add_bullets(doc, [
        "所有方块后续都使用单独 class，注册类只负责接线。",
        "新增物品必须有贴图、模型、语言和至少一个真实用途。",
        "机器系统围绕 Forge Energy、物品处理、可升级自动化逐步展开。",
        "开发节奏按阶段推进，避免出现大量紫黑贴图或没有功能的注册项。"
    ])

    add_heading(doc, "2. 核心时代线", 1)
    table = doc.add_table(rows=1, cols=4)
    style_table(table)
    headers = ["阶段", "主题", "核心内容", "进入下一阶段条件"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    rows = [
        ("阶段 1", "手工科技", "科技工作台、基础零件、专属配方入口", "工作台 UI、基础配方、贴图全部稳定"),
        ("阶段 2", "燃煤/蒸汽", "燃煤发电机、基础粉碎机、电炉", "形成发电到加工的最小闭环"),
        ("阶段 3", "电力时代", "线缆、电池、储能方块、机器能量槽", "机器能相邻供电并显示能量"),
        ("阶段 4", "自动化", "管道、过滤器、自动输入输出、红石控制", "能搭建自动矿物处理线"),
        ("阶段 5", "高级材料", "合金炉、压缩机、硅片、电路等级", "高级机器需要多步骤材料链"),
        ("阶段 6", "高能科技", "太阳能阵列、反应堆、量子核心、终局制造", "玩家拥有长期追求目标"),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_text(cells[i], value)

    add_heading(doc, "3. 机器体系", 1)
    add_heading(doc, "3.1 基础机器", 2)
    add_bullets(doc, [
        "燃煤发电机：燃料转 FE，早期电力来源。",
        "矿石粉碎机：矿石或原矿变粉末，提供双倍矿物处理入口。",
        "电力熔炉：消耗 FE 完成烧炼，比原版熔炉更快。",
        "压缩机：粉末、板材、晶圆等中级材料加工。",
        "合金炉：钢、青铜、电工钢、钛合金等高级材料来源。"
    ])
    add_heading(doc, "3.2 高级机器", 2)
    add_bullets(doc, [
        "电路打印机：制造基础/高级/精密电路板。",
        "自动装配机：按模板自动合成科技配方。",
        "采矿钻机：消耗大量能量自动采集矿物。",
        "化学反应器：处理液体、晶体、硅片和高级材料。",
        "多方块反应堆：高风险高产出的终局能源系统。"
    ])

    add_heading(doc, "4. 能源系统", 1)
    add_paragraph(doc, "能源使用 Forge Energy。第一版不直接做复杂电网，先让每台机器拥有自己的能量存储，并支持相邻方块供电。线缆、电池、储能阵列在后续阶段逐步加入。")
    add_bullets(doc, [
        "基础线缆：便宜、传输低。",
        "铜线缆：中等传输，作为早期主力。",
        "金线缆：高传输，成本更高。",
        "超导线缆：终局传输，无损耗或极低损耗。",
        "储能单元：从小电池盒升级到大型储能阵列。"
    ])

    add_heading(doc, "5. 材料与配方链", 1)
    add_paragraph(doc, "材料链要服务玩法，而不是单纯增加物品数量。每个材料都应连接到机器、升级、线缆或阶段解锁。")
    add_bullets(doc, [
        "基础材料：铜线、铁板、齿轮、硅片、基础电路板。",
        "矿物材料：锡、铝、铅、银、钛、铀。",
        "加工链：矿石 -> 粉碎 -> 粉末 -> 洗矿 -> 精矿 -> 熔炼。",
        "副产物：粉碎铜矿可少量获得金粉，粉碎铁矿可少量获得镍粉。",
        "合金：钢、青铜、电工钢、钛合金。"
    ])

    add_heading(doc, "6. 自动化系统", 1)
    add_bullets(doc, [
        "机器六面输入/输出配置。",
        "物品管道支持抽取、插入、过滤和优先级。",
        "红石控制支持始终运行、红石启用、红石禁用。",
        "机器升级支持速度、能耗、容量、并行处理。",
        "物流网络后期支持自动合成和库存请求。"
    ])

    add_heading(doc, "7. 第一阶段：科技工作台", 1)
    add_paragraph(doc, "科技工作台是模组入口。第一阶段不追求复杂系统，而是把基础结构搭好：独立方块类、方块实体、菜单、客户端屏幕、贴图、语言和可堆叠可放置的方块物品。")
    add_numbered(doc, [
        "TechWorkbenchBlock：负责方块交互、打开界面、创建方块实体。",
        "TechWorkbenchBlockEntity：保存方块状态，并作为菜单提供者。",
        "TechWorkbenchMenu：服务端容器结构，先接入玩家背包，后续加入专属合成槽。",
        "TechWorkbenchScreen：客户端界面，后续替换成纹理 GUI 并显示专属槽位。",
        "注册规范：TechMod 只注册方块、物品、方块实体、菜单和创造栏，不把行为堆在注册类里。"
    ])

    add_heading(doc, "8. 开发规范", 1)
    add_bullets(doc, [
        "每个方块单独 class，后续机器方块也必须如此。",
        "每个机器至少包含 block、block entity、menu、screen 四层结构。",
        "每个新增资源必须同步 blockstate、block model、item model、texture、lang。",
        "新增注册项前先做贴图和模型，禁止先注册一堆不可见内容。",
        "每个阶段完成后再进入下一阶段，避免系统半成品互相污染。"
    ])

    add_heading(doc, "9. 短期任务清单", 1)
    add_bullets(doc, [
        "完成科技工作台自定义 GUI 的基础显示。",
        "加入工作台专属 3x3 或扩展合成槽。",
        "新增第一批真实零件：铜线、铁板、齿轮、基础电路板、硅片。",
        "为每个零件制作贴图、模型、语言和配方。",
        "准备燃煤发电机的方块类、方块实体和能量接口设计。"
    ])

    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
